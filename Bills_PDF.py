"""对账单生成"""
import os

import pandas as pd
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from win32com.client import constants, gencache
from utilities import get_qrcode, get_barcode
from docx.shared import Inches
from win32com import client

OutputFolder = 'D:/Code/python/temp/'
qrCodeOutPutFolder = OutputFolder + 'qrcode/'
barCodeOutPutFolder = OutputFolder + 'barcode/'


def generate_word(company_filename, template_filename):
    # 读取Excel文件
    excel = pd.read_excel(company_filename, sheet_name='天职制函表')
    # 遍历excel每一行
    for index, row in excel.iterrows():
        # 通过列名访问单元格数据
        name = row['客户名称']
        amount = "{:,.2f}".format(row['(函证)交易额（1-6月）'])
        ying_shou = "{:,.2f}".format(row['(函证)应收余额'])
        yu_shou = "{:,.2f}".format(row['(函证)预收余额'])
        fang_li = "{:,.2f}".format(row['(发函)其中：返利'])
        org_name = row['索引']
        shipping_code = row['快递备注']
        unique_id = row['唯一识别码']

        doc = Document(template_filename)

        read_document(doc, name)
        read_tables(doc, amount, ying_shou, yu_shou, fang_li, org_name, shipping_code)
        unique_id_bar_code(doc, unique_id)
        add_qr_code(doc, shipping_code)
        add_bar_code(doc, shipping_code)

        doc.save(OutputFolder + name + '.docx')
        print('生成：' + OutputFolder + name + '.docx')

        word_to_pdf_new(OutputFolder + name + '.docx', OutputFolder + "pdf/" + name + ".pdf")
        print('生成：' + OutputFolder + "pdf/" + name + ".pdf")


def unique_id_bar_code(document, unique_id):

    # 唯一识别码
    bar_code_cell = document.tables[5].cell(1, 1)
    bar_img_path = barCodeOutPutFolder + unique_id
    get_barcode(unique_id, bar_img_path)
    bar_code_run = bar_code_cell.paragraphs[0].add_run()
    bar_code_run.add_picture(bar_img_path + ".png", width=Inches(2.5), height=Inches(0.8))


def add_bar_code(document, bar_code):

    # 条形码
    bar_code_cell = document.tables[4].cell(0, 0)
    bar_img_path = barCodeOutPutFolder + bar_code
    get_barcode(bar_code, bar_img_path)
    bar_code_run = bar_code_cell.paragraphs[0].add_run()
    bar_code_run.add_picture(bar_img_path + ".png", width=Inches(2.5), height=Inches(0.8))


def add_qr_code(document, code):
    # 二维码
    cell = document.tables[1].cell(2, 4)
    img_path = qrCodeOutPutFolder + code + ".jpg"
    get_qrcode(code, img_path)
    run = cell.paragraphs[0].add_run()
    run.add_picture(img_path, width=Inches(0.8))


def read_document(document, name):
    # 遍历文档
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            # 替换功能
            if "«client»" in run.text:
                run.text = run.text.replace("«client»", " " + name)


def read_tables(document, amount, ying_shou, yu_shou, fang_li, org_name, shipping_code):
    # 遍历表格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # 遍历表格段落内容，回到上个步骤，将cell当作paragraph处理
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # 替换功能
                        if "«jiaoyie»" in cell.text:
                            run.text = run.text.replace("«jiaoyie»", amount)
                            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "«yingszk»" in cell.text:
                            run.text = run.text.replace("«yingszk»", ying_shou)
                        if "«yuszk»" in cell.text:
                            run.text = run.text.replace("«yuszk»", yu_shou)
                        if "«fanli»" in cell.text:
                            run.text = run.text.replace("«fanli»", fang_li)
                        if "waybillnumber" in cell.text:
                            run.text = run.text.replace("waybillnumber", shipping_code)
                        if "«unique_code»" in cell.text:
                            run.text = run.text.replace("«unique_code»", ' ')
                        if "«org_name»" in cell.text:
                            run.text = run.text.replace("«org_name»", org_name)


def get_file_names(path, filetype):
    file_names = []
    for file in os.listdir(path):
        if os.path.isfile(os.path.join(path, file)) and file.find(filetype) > -1:
            file_names.append(path + file)
    return file_names


def word_to_pdf_new(word_path, save_path):
    word = client.Dispatch("Word.Application")
    doc = word.Documents.Open(word_path)
    doc.SaveAs(save_path, 17)
    doc.Close()
    word.Quit()


def merge_pdf(folder_name, merge_pdf_name):
    pdf_files = get_file_names(folder_name, 'pdf')
    pdf_writer = PdfWriter()
    for pdf_file in pdf_files:
        pdf_reader = PdfReader(pdf_file)
        for page in range(len(pdf_reader.pages)):
            # Add each page to the writer object
            pdf_writer.add_page(pdf_reader.pages[page])
    # Write out the merged PDF
    with open(folder_name + merge_pdf_name, 'wb') as out:
        pdf_writer.write(out)
